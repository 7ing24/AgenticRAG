"""结构化文档提取：PDF/docx 的正文 + 表格（转 Markdown）

- PDF 用 pdfplumber：逐页提取正文，表格区域挖洞后转 Markdown
- docx 用 python-docx：按文档顺序遍历段落与表格

表格统一转成 Markdown 表格，保留行列结构，供下游切分 + 向量化。
"""

import logging
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _clean_cell(text) -> str:
    """清洗单元格：换行→空格、转义管道符"""
    if text is None:
        return ""
    text = str(text)
    text = text.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    text = text.replace("|", "\\|")
    return text.strip()


def _table_to_markdown(rows: List[List]) -> str:
    """把二维表格转成 Markdown 表格文本"""
    if not rows:
        return ""

    # 过滤全空行
    rows = [r for r in rows if r is not None and any((c or "").strip() for c in r)]
    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    # 规范化：补齐到最大列数
    norm = []
    for r in rows:
        row = list(r) + [""] * (max_cols - len(r))
        norm.append(row[:max_cols])

    lines = []
    header = norm[0]
    lines.append("| " + " | ".join(_clean_cell(c) for c in header) + " |")
    lines.append("|" + " --- |" * max_cols)
    for r in norm[1:]:
        lines.append("| " + " | ".join(_clean_cell(c) for c in r) + " |")
    return "\n".join(lines)


def _extract_pdf_page(page) -> str:
    """提取单页：正文挖掉表格区域 + 表格 Markdown（按 top 排序附后）"""
    tables = page.find_tables() or []
    if not tables:
        return page.extract_text() or ""

    # 表格转 markdown，记录 top 坐标（bbox = (x0, top, x1, bottom)）
    table_blocks = []
    for t in tables:
        md = _table_to_markdown(t.extract() or [])
        if md:
            table_blocks.append((t.bbox[1], md))

    # 挖掉表格区域，提取正文（避免表格文字在正文里重复）
    body = page
    for t in tables:
        try:
            body = body.outside_bbox(t.bbox)
        except Exception:
            pass
    text = body.extract_text() or ""

    # 拼接：正文在前 + 各表按 top 顺序附后（折中：不重建像素级阅读顺序）
    table_blocks.sort(key=lambda b: b[0])
    parts = [text] + [md for _, md in table_blocks]
    return "\n\n".join(p for p in parts if p and p.strip())


def extract_pdf(file_path: str) -> List[Document]:
    """PDF 提取：正文 + Markdown 表格，每页一个 Document"""
    import pdfplumber

    documents = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                content = _extract_pdf_page(page)
                if not content or not content.strip():
                    continue
                documents.append(Document(
                    page_content=content,
                    metadata={"source": file_path, "page": page_idx + 1},
                ))
    except Exception as e:
        logger.error(f"[TableExtractor] pdfplumber failed to parse PDF: {e}")
        raise
    return documents


def extract_docx(file_path: str) -> List[Document]:
    """docx 提取：按文档顺序遍历段落与表格，单个 Document（兼容 Docx2txtLoader）"""
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    doc = DocxDocument(file_path)
    parts = []

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            md = _table_to_markdown(rows)
            if md:
                parts.append(md)

    content = "\n\n".join(parts)
    if not content.strip():
        return []
    return [Document(page_content=content, metadata={"source": file_path, "page": 1})]


def extract_html(file_path: str) -> List[Document]:
    """HTML 提取：去脚本/样式/导航，表格转 Markdown，提取正文"""
    from bs4 import BeautifulSoup

    with open(file_path, "rb") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # 去掉脚本、样式、导航等噪音标签
    for tag in soup(["script", "style", "noscript", "nav", "header", "footer", "aside", "iframe"]):
        tag.decompose()

    # 表格转 Markdown 后从 DOM 移除（避免 get_text 重复提取）
    table_mds = []
    for table in soup.find_all("table"):
        rows = [
            [td.get_text(" ", strip=True) for td in tr.find_all(["th", "td"])]
            for tr in table.find_all("tr")
        ]
        md = _table_to_markdown(rows)
        if md:
            table_mds.append(md)
        table.decompose()

    text = (soup.body or soup).get_text("\n")
    parts = [text] + table_mds
    content = "\n\n".join(p.strip() for p in parts if p and p.strip())
    if not content:
        return []
    return [Document(page_content=content, metadata={"source": file_path, "page": 1})]
